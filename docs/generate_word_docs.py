"""
Generate Word documents:
1. PROZESSAENDERUNGEN_V2.docx - Full process document
2. IST_SOLL_VERGLEICH.docx - Side-by-side comparison table
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elem)


def set_cell_text(cell, text, bold=False, size=9, color=None, align=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_header_row(table, texts, color_hex="2B579A"):
    """Format the first row of a table as header."""
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        set_cell_shading(cell, color_hex)
        set_cell_text(cell, text, bold=True, size=10, color=(255, 255, 255))


def add_data_row(table, texts, change_type=None):
    """Add a data row to a table with optional color coding."""
    row = table.add_row()
    for i, text in enumerate(texts):
        set_cell_text(row.cells[i], text, size=9)

    # Color the last column based on change type
    if change_type:
        last_cell = row.cells[len(texts) - 1]
        colors = {
            "GEÄNDERT": "FFF2CC",
            "NEU": "D9EAD3",
            "ENTFÄLLT": "F4CCCC",
            "ERWEITERT": "D0E0F0",
            "VEREINFACHT": "E8D5F5",
            "PRÄZISIERT": "D0E0F0",
            "UNVERÄNDERT": "EEEEEE",
        }
        bg = colors.get(change_type, "FFFFFF")
        set_cell_shading(last_cell, bg)
        set_cell_text(last_cell, change_type, bold=True, size=9)

    return row


# ══════════════════════════════════════════════════════════════════════════
# DOCUMENT 1: Full PROZESSAENDERUNGEN_V2.docx
# ══════════════════════════════════════════════════════════════════════════

def create_prozessaenderungen():
    doc = Document()

    # Page setup: A4 landscape for wide tables
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Title ─────────────────────────────────────────────────────────
    title = doc.add_heading("Prozessvergleich & Änderungen V2", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Stand: 03.03.2026 | CSFV Minden e.V. — Betreuer-App")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # ── 1. REGISTRIERUNGSPROZESS ──────────────────────────────────────
    doc.add_heading("1. Registrierungsprozess — Betreuer", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["#", "AKTUELL (IST)", "NEU (SOLL)", "Status"])

    reg_rows = [
        ("1.1", "Koordinator erstellt token-basierten Registrierungslink (single/multi-use, mit Ablaufdatum)", "Fester, permanenter Registrierungslink — immer erreichbar, kein Token nötig", "GEÄNDERT"),
        ("1.2", "Betreuer öffnet /registrierung/<token>/", "Betreuer öffnet /registrierung/ (feste URL)", "GEÄNDERT"),
        ("1.3", "Betreuer füllt persönliche Daten aus (Name, Adresse, Geburtsdatum, Bank)", "Betreuer füllt persönliche Daten aus + wählt Tätigkeitsart + wählt Schule", "ERWEITERT"),
        ("1.4", "System erstellt User-Account mit generiertem Passwort", "Betreuer legt eigenes Passwort fest im Registrierungsformular", "GEÄNDERT"),
        ("1.5", "System prüft nicht auf Duplikate", "Eindeutigkeitsprüfung per Hash (z.B. Vorname+Nachname+Geburtsdatum). Duplikate werden erkannt", "NEU"),
        ("1.6", "Keine Email-Prüfung gegen Bestand", "Email-Abgleich: Wenn Betreuer sich mit anderer Email registriert als im System bekannt → Hinweis", "NEU"),
        ("1.7", "Jeder Betreuer hat genau einen Vertrag pro Schuljahr", "Betreuer muss sich für jedes Förderprogramm/Schule einzeln anmelden — mehrere Anmeldungen möglich", "GEÄNDERT"),
        ("1.8", "Führungszeugnis nur für externe Betreuer (is_external=True)", "Führungszeugnis nur wenn Betreuer volljährig (>= 18 Jahre). Minderjährige brauchen keins", "GEÄNDERT"),
        ("1.9", "IBAN wird mit Fernet verschlüsselt gespeichert", "IBAN wird unverschlüsselt gespeichert (kein EncryptedCharField)", "VEREINFACHT"),
        ("1.10", "Email nicht als Pflichtfeld", "Jeder Betreuer muss eine Emailadresse haben (Pflichtfeld)", "NEU"),
    ]
    for nr, ist, soll, status in reg_rows:
        add_data_row(table, [nr, ist, soll, status], status)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(10.5)
        row.cells[2].width = Cm(11.5)
        row.cells[3].width = Cm(3.0)

    doc.add_paragraph()

    # ── Tätigkeitsarten ───────────────────────────────────────────────
    doc.add_heading("Tätigkeitsarten bei der Registrierung (NEU)", level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Code", "Tätigkeitsart", "Verfügbar in Förderprogramm"])

    taet_rows = [
        ("ag_leitung", "AG-Leitung", "Geld oder Stelle, 13 plus"),
        ("hausaufgabenbetreuung", "Hausaufgabenbetreuung", "Geld oder Stelle, Schule von 8 bis 1"),
        ("hausaufgabenhilfe_plus", "Hausaufgabenhilfe plus", "Geld oder Stelle"),
        ("aufsicht", "Aufsicht", "Schule von 8 bis 1"),
        ("paed_assistenz", "Pädagogische Assistenz", "Schule von 8 bis 1"),
        ("schwimmbegleitung", "Schwimmbegleitung", "Geld oder Stelle"),
    ]
    for code, name, fp in taet_rows:
        add_data_row(table, [code, name, fp])

    doc.add_paragraph()

    # ── Schulauswahl ──────────────────────────────────────────────────
    doc.add_heading("Schulauswahl & Standard-Förderprogramm", level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Schule", "Schultyp", "Standard-Förderprogramm"])

    school_rows = [
        ("Gesamtschule", "weiterfuehrend", "Geld oder Stelle"),
        ("Gymnasium", "weiterfuehrend", "Geld oder Stelle"),
        ("Grundschule Haddenhausen", "grundschule", "Schule von 8 bis 1"),
        ("Grundschule Minderheide", "grundschule", "Schule von 8 bis 1"),
        ("Grundschule Stemwede", "grundschule", "Schule von 8 bis 1"),
    ]
    for s, t, fp in school_rows:
        add_data_row(table, [s, t, fp])

    p = doc.add_paragraph()
    run = p.add_run("Sonderregel: ")
    run.bold = True
    p.add_run("AG-Leitung + Grundschule → Förderprogramm = \"13 plus\" (nicht Schule von 8 bis 1)")

    doc.add_paragraph()

    # ── Standard-Tätigkeiten je FP ────────────────────────────────────
    doc.add_heading("Standard-Tätigkeiten je Förderprogramm", level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Förderprogramm", "Schulkategorie", "Standard-Tätigkeiten"])

    fp_rows = [
        ("Geld oder Stelle", "Weiterführend (GES/GYM)", "AG-Leitung, Hausaufgabenbetreuung, Hausaufgabenhilfe plus, Schwimmbegleitung"),
        ("Schule von 8 bis 1", "Grundschule", "Hausaufgabenbetreuung, Päd. Assistenz, Aufsicht"),
        ("13 plus", "Grundschule (nur AG)", "AG-Leitung"),
    ]
    for fp, kat, taet in fp_rows:
        add_data_row(table, [fp, kat, taet])

    doc.add_page_break()

    # ── 2. GENEHMIGUNGSPROZESS ────────────────────────────────────────
    doc.add_heading("2. Genehmigungsprozess — Koordinator", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["#", "AKTUELL (IST)", "NEU (SOLL)", "Status"])

    koord_rows = [
        ("2.1", "Koordinator erstellt Registrierungslinks", "Entfällt — Registrierungslink ist fest/permanent", "ENTFÄLLT"),
        ("2.2", "Koordinator sieht eingegangene Registrierungen für seine Schule(n)", "Unverändert: Koordinator sieht Anmeldungen für seine zugewiesene(n) Schule(n)", "UNVERÄNDERT"),
        ("2.3", "Koordinator genehmigt/lehnt ab", "Unverändert: Koordinator genehmigt jede Anmeldung", "UNVERÄNDERT"),
        ("2.4a", "—", "NEU: Koordinator ergänzt Förderprogramm (Standard wird vorgeschlagen)", "NEU"),
        ("2.4b", "—", "NEU: Koordinator legt Vertragsbeginn fest", "NEU"),
        ("2.4c", "—", "NEU: Koordinator ergänzt Betreuer-Typ (Schüler, sonst. Mitarbeiter, langjährig, Lehrer, LA-Student, extern)", "NEU"),
        ("2.4d", "—", "NEU: Bei AG-Leitung → Koordinator trägt genauen AG-Namen ein", "NEU"),
        ("2.4e", "—", "NEU: Koordinator ordnet/bestätigt Tätigkeit zu", "NEU"),
        ("2.5", "Nach Genehmigung: Dokumente werden generiert", "Unverändert: Email an Betreuer + Email an Personalabteilung (via N8N)", "UNVERÄNDERT"),
        ("2.6", "Koordinator genehmigt Stundennachweise", "Unverändert: Koordinator genehmigt Stundennachweise für seine Schule(n)", "UNVERÄNDERT"),
    ]
    for nr, ist, soll, status in koord_rows:
        add_data_row(table, [nr, ist, soll, status], status)

    for row in table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(10.5)
        row.cells[2].width = Cm(11.5)
        row.cells[3].width = Cm(3.0)

    doc.add_paragraph()

    # ── 3. ADMIN ──────────────────────────────────────────────────────
    doc.add_heading("3. Admin-Prozesse", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["#", "AKTUELL (IST)", "NEU (SOLL)", "Status"])

    admin_rows = [
        ("3.1", "Admin verwaltet Schulen, Schuljahre, Förderprogramme, Kostenstellen", "Unverändert", "UNVERÄNDERT"),
        ("3.2", "Admin verwaltet Stundensätze (HourlyRate)", "Unverändert", "UNVERÄNDERT"),
        ("3.3", "Admin aktiviert Betreuer (Status → aktiv)", "Unverändert", "UNVERÄNDERT"),
        ("3.4", "Freibetrag-Limit als Feld in SchoolYear (3.300 EUR)", "§3 Nr. 26 EStG Übungsleiterpauschale als eigenes Stammdatum — pro Kalenderjahr konfigurierbar", "GEÄNDERT"),
        ("3.5", "—", "NEU: Manuelle Kostenbelastung pro Förderprogramm (Admin kann manuell Kosten zuordnen)", "NEU"),
        ("3.6", "Keine Auswertung pro Förderprogramm", "NEU: Auswertungen pro Förderprogramm — tatsächlich gezahlte Beträge (über Kostenstelle)", "NEU"),
        ("3.7", "Admin verwaltet Registrierungslinks", "Entfällt — Link ist permanent", "ENTFÄLLT"),
    ]
    for nr, ist, soll, status in admin_rows:
        add_data_row(table, [nr, ist, soll, status], status)

    for row in table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(10.5)
        row.cells[2].width = Cm(11.5)
        row.cells[3].width = Cm(3.0)

    doc.add_page_break()

    # ── 4. ZEITERFASSUNG ──────────────────────────────────────────────
    doc.add_heading("4. Zeiterfassung & Abrechnung", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["#", "AKTUELL (IST)", "NEU (SOLL)", "Status"])

    zeit_rows = [
        ("4.1", "Betreuer erfasst Zeiten pro Vertrag", "Betreuer erfasst Zeiten separat pro Schule", "GEÄNDERT"),
        ("4.2", "Alle Zeiten gehen an einen Koordinator", "Je nach Schule geht die Genehmigung an den jeweiligen Koordinator", "GEÄNDERT"),
        ("4.3", "Ein Vertrag pro Förderprogramm", "Betreuer kann an verschiedenen Förderprogrammen pro Schule angemeldet sein", "ERWEITERT"),
        ("4.4", "Zahlungsanweisung nach Genehmigung", "Zahlungsanweisung wird pro Förderprogramm erstellt (Kreditorennr., Projektnr., Betrag kodiert)", "PRÄZISIERT"),
        ("4.5", "Koordinator sieht Stundennachweise seiner Schule(n)", "Unverändert", "UNVERÄNDERT"),
        ("4.6", "Genehmigung → PDF → N8N-Benachrichtigung", "Unverändert: Zahlungsanweisung pro Förderprogramm mit QR-Code", "UNVERÄNDERT"),
        ("4.7", "Ablehnung → Betreuer korrigiert & reicht erneut ein", "Unverändert", "UNVERÄNDERT"),
    ]
    for nr, ist, soll, status in zeit_rows:
        add_data_row(table, [nr, ist, soll, status], status)

    for row in table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(10.5)
        row.cells[2].width = Cm(11.5)
        row.cells[3].width = Cm(3.0)

    doc.add_paragraph()

    # ── 5. BUDGET & FREIBETRAG ────────────────────────────────────────
    doc.add_heading("5. Budget & Freibetrag", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["#", "AKTUELL (IST)", "NEU (SOLL)", "Status"])

    budget_rows = [
        ("5.1", "Freibetrag-Limit als Feld in SchoolYear (default 3.300 EUR)", "§3 Nr. 26 EStG Übungsleiterpauschale als eigenes Stammdatum — pro Kalenderjahr", "GEÄNDERT"),
        ("5.2", "Freibetrag ist Kalenderjahr-bezogen (01.01–31.12)", "Unverändert", "UNVERÄNDERT"),
        ("5.3", "Budget pro Förderprogramm ist Schuljahr-bezogen", "Unverändert", "UNVERÄNDERT"),
        ("5.4", "Auszahlungen nicht explizit Kalenderjahr-bezogen", "Auszahlungen sind Kalenderjahr-bezogen — Budget an §3 Nr. 26 EStG-Betrag geknüpft", "PRÄZISIERT"),
        ("5.5", "Kostenstelle hat FK zu Förderprogramm", "Erweitert: Auswertung pro Förderprogramm über Kostenstelle möglich", "ERWEITERT"),
    ]
    for nr, ist, soll, status in budget_rows:
        add_data_row(table, [nr, ist, soll, status], status)

    for row in table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(10.5)
        row.cells[2].width = Cm(11.5)
        row.cells[3].width = Cm(3.0)

    doc.add_paragraph()

    # ── 6. DOKUMENTEN-PROZESS ─────────────────────────────────────────
    doc.add_heading("6. Dokumenten-Prozess", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["#", "AKTUELL (IST)", "NEU (SOLL)", "Status"])

    doc_rows = [
        ("6.1", "Führungszeugnis nötig wenn is_external=True", "Führungszeugnis nötig wenn Betreuer volljährig (>= 18 Jahre)", "GEÄNDERT"),
        ("6.2", "Vertrag wird auto-generiert nach Registrierung", "Vertrag wird generiert nach Genehmigung durch Koordinator (da Koordinator Vertragsbeginn + FP ergänzt)", "GEÄNDERT"),
        ("6.3", "Vertraulichkeitserklärung auto-generiert", "Unverändert", "UNVERÄNDERT"),
        ("6.4", "IfSB auto-generiert (24 Monate Renewal)", "Unverändert", "UNVERÄNDERT"),
        ("6.5", "Masernschutz manueller Upload", "Unverändert", "UNVERÄNDERT"),
    ]
    for nr, ist, soll, status in doc_rows:
        add_data_row(table, [nr, ist, soll, status], status)

    for row in table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(10.5)
        row.cells[2].width = Cm(11.5)
        row.cells[3].width = Cm(3.0)

    doc.add_page_break()

    # ── 7. DATENMODELL-ÄNDERUNGEN ─────────────────────────────────────
    doc.add_heading("7. Datenmodell-Änderungen", level=1)

    doc.add_heading("7.1 RegistrationLink → ENTFÄLLT", level=2)
    doc.add_paragraph("Token-basierte Links werden durch einen festen Registrierungsendpunkt ersetzt.")

    doc.add_heading("7.2 BetreuerProfile → Anpassungen", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Feld", "IST", "SOLL"])
    model_rows = [
        ("iban", "EncryptedCharField (Fernet-verschlüsselt)", "Normales CharField (unverschlüsselt)"),
        ("unique_hash", "— (nicht vorhanden)", "NEU: Hash aus Vorname+Nachname+Geburtsdatum"),
        ("email (User)", "Optional", "Pflichtfeld, muss validiert werden"),
        ("requires_fuehrungszeugnis", "Basiert auf is_external", "Basiert auf Alter >= 18 bei Anmeldung"),
    ]
    for feld, ist, soll in model_rows:
        add_data_row(table, [feld, ist, soll])

    doc.add_paragraph()
    doc.add_heading("7.3 Contract → Anpassungen", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Feld", "IST", "SOLL"])
    contract_rows = [
        ("start_date", "Automatisch bei Registrierung", "Vom Koordinator bei Genehmigung gesetzt"),
        ("foerderprogramme", "Optional, manuell zugeordnet", "Vom Koordinator bei Genehmigung zugeordnet (Standard vorgeschlagen)"),
    ]
    for feld, ist, soll in contract_rows:
        add_data_row(table, [feld, ist, soll])

    doc.add_paragraph()
    doc.add_heading("7.4 ActivityType → Aktualisierte Liste", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    add_header_row(table, ["Code", "Tätigkeitsart"])
    for code, name in [("ag_leitung", "AG-Leitung"), ("hausaufgabenbetreuung", "Hausaufgabenbetreuung"),
                        ("hausaufgabenhilfe_plus", "Hausaufgabenhilfe plus"), ("aufsicht", "Aufsicht"),
                        ("paed_assistenz", "Pädagogische Assistenz"), ("schwimmbegleitung", "Schwimmbegleitung")]:
        add_data_row(table, [code, name])

    doc.add_paragraph()
    doc.add_heading("7.5 Übungsleiterpauschale → NEUES MODELL", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    add_header_row(table, ["Feld", "Beschreibung"])
    for feld, desc in [("kalenderjahr", "z.B. 2026"), ("betrag", "z.B. 3.300,00 EUR"),
                        ("gesetzliche_grundlage", "§3 Nr. 26 EStG")]:
        add_data_row(table, [feld, desc])

    doc.add_paragraph()
    doc.add_heading("7.6 ManuelleKostenbuchung → NEUES MODELL", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    add_header_row(table, ["Feld", "Beschreibung"])
    for feld, desc in [("foerderprogramm", "FK zu Förderprogramm"), ("betrag", "Manuell eingetragener Betrag"),
                        ("beschreibung", "Freitext"), ("datum", "Buchungsdatum"),
                        ("erstellt_von", "FK zu User (Admin)")]:
        add_data_row(table, [feld, desc])

    doc.add_page_break()

    # ── 8. FLUSSDIAGRAMME ─────────────────────────────────────────────
    doc.add_heading("8. Prozess-Flussdiagramme", level=1)

    doc.add_heading("8.1 Registrierung & Genehmigung", level=2)

    # Flow as a table (since Word doesn't do ASCII art well)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Schritt", "Wer", "Aktion"])

    flow_reg = [
        ("1", "Betreuer", "Öffnet festen Registrierungslink /registrierung/"),
        ("2", "Betreuer", "Füllt Formular aus: Persönliche Daten, Adresse, Bank, Passwort, Tätigkeitsart, Schule"),
        ("3", "System", "Prüft Hash-Duplikat und Email-Abgleich"),
        ("4", "System", "Speichert Registrierung (Status: registered)"),
        ("5", "Koordinator", "Sieht neue Anmeldung für seine Schule"),
        ("6", "Koordinator", "Ergänzt: Förderprogramm, Vertragsbeginn, Betreuer-Typ, ggf. AG-Name, Tätigkeit"),
        ("7", "Koordinator", "Genehmigt die Anmeldung ✓"),
        ("8", "System", "Email an Betreuer (Genehmigung) + Email an Personalabteilung (via N8N)"),
        ("9", "System", "Generiert Dokumente: Vertrag, Vertraulichkeitserklärung, IfSB, ggf. Führungszeugnis"),
        ("10", "Betreuer", "Lädt signierte Dokumente hoch"),
        ("11", "Koordinator", "Prüft & verifiziert Dokumente"),
        ("12", "Admin", "Aktiviert Betreuer (Status → aktiv)"),
    ]
    for schritt, wer, aktion in flow_reg:
        row = add_data_row(table, [schritt, wer, aktion])
        # Color code by role
        colors = {"Betreuer": "D0E0F0", "Koordinator": "D9EAD3", "System": "EEEEEE", "Admin": "FFF2CC"}
        set_cell_shading(row.cells[1], colors.get(wer, "FFFFFF"))

    doc.add_paragraph()
    doc.add_heading("8.2 Zeiterfassung & Abrechnung", level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Schritt", "Wer", "Aktion"])

    flow_zeit = [
        ("1", "Betreuer", "Wählt Schule aus (wenn mehrere)"),
        ("2", "Betreuer", "Erfasst Zeiten: Datum, Start/Ende, Pause, Förderprogramm, Beschreibung"),
        ("3", "Betreuer", "Reicht Monat ein (Status: submitted)"),
        ("4", "Koordinator", "Sieht Stundennachweis für seine Schule"),
        ("5a", "Koordinator", "Genehmigt ✓ → PDF-Zahlungsanweisung pro FP, N8N-Benachrichtigung, Freibetrag-Prüfung"),
        ("5b", "Koordinator", "ODER: Lehnt ab ✗ → Betreuer korrigiert & reicht erneut ein"),
    ]
    for schritt, wer, aktion in flow_zeit:
        row = add_data_row(table, [schritt, wer, aktion])
        colors = {"Betreuer": "D0E0F0", "Koordinator": "D9EAD3"}
        set_cell_shading(row.cells[1], colors.get(wer, "FFFFFF"))

    doc.add_page_break()

    # ── 9. OFFENE RÜCKFRAGEN ──────────────────────────────────────────
    doc.add_heading("9. Offene Rückfragen", level=1)

    questions = [
        ("Frage 1: Formulare", "Können die Links zu den zwei Formularen (FES-Personal-Betreuerfragebogen Betreuer + Koordinator) geteilt werden, um die genauen Felder ins System zu übernehmen?"),
        ("Frage 2: Hash-Bildung", "Aus welchen Feldern wird der Hash gebildet? Vorschlag: Vorname + Nachname + Geburtsdatum → SHA256-Hash"),
        ("Frage 3: Email-Abgleich", "Gegen welche Daten wird abgeglichen? (a) Importierte Liste aus altem System? (b) Bereits registrierte Betreuer (Hash-Match)?"),
        ("Frage 4: Mehrfach-Anmeldung", "Voller Registrierungsprozess nochmal oder vereinfacht (persönliche Daten wiederverwenden)?"),
        ("Frage 5: Gesamtschule & Gymnasium", "Zwei separate Schulen mit eigenem Koordinator oder eine kombinierte Option?"),
        ("Frage 6: Vertragsende", "Automatisch auf Schuljahresende oder manuell vom Koordinator?"),
        ("Frage 7: Onboarding-Status", "Neuer Status 'pending_approval' nötig zwischen Registrierung und Koordinator-Genehmigung?"),
    ]
    for title_text, desc in questions:
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ── 10. VERBESSERUNGSVORSCHLÄGE ───────────────────────────────────
    doc.add_heading("10. Verbesserungsvorschläge", level=1)

    suggestions = [
        ("Vorschlag 1: Intelligente Formular-Vorauswahl",
         "Bei der Registrierung basierend auf Schulauswahl automatisch nur passende Tätigkeitsarten anzeigen. "
         "Grundschule → Hausaufgabenbetreuung, Päd. Assistenz, Aufsicht, AG-Leitung. "
         "Gesamtschule/Gymnasium → AG-Leitung, Hausaufgabenbetreuung, Hausaufgabenhilfe plus, Schwimmbegleitung."),
        ("Vorschlag 2: Koordinator-Standardwerte",
         "Standardwerte vorbelegen: Förderprogramm automatisch vorschlagen basierend auf Schule + Tätigkeit. "
         "Vertragsbeginn → nächster Monatserster als Default."),
        ("Vorschlag 3: Dashboard für Mehrfach-Schulen",
         "Betreuer-Dashboard: Zeiten nach Schule gruppiert, Freibetrag über alle Schulen aggregiert, "
         "Verträge pro Schule/Förderprogramm auflisten."),
        ("Vorschlag 4: Automatische Förderprogramm-Zuordnung",
         "Deterministische Zuordnungslogik: Grundschule + AG-Leitung → 13 plus, "
         "Grundschule + andere → Schule von 8 bis 1, Weiterführend → Geld oder Stelle. "
         "Koordinator bestätigt nur noch statt manuell zuzuordnen."),
        ("Vorschlag 5: Übungsleiterpauschale als eigenes Modell",
         "Eigenes Modell mit Kalenderjahr + Betrag + gesetzliche Grundlage. "
         "Trennt sauber Schuljahr-Budget von Kalenderjahr-Freibetrag."),
        ("Vorschlag 6: Manuelle Kostenbuchung mit Kategorisierung",
         "Kategorien: Material, Fortbildung, Versicherung, Sonstiges. "
         "Zusätzlich: Beleg-Nr. für Buchhaltung."),
    ]
    for title_text, desc in suggestions:
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph(desc)

    doc.add_paragraph()

    # ── 11. ZUSAMMENFASSUNG ───────────────────────────────────────────
    doc.add_heading("11. Zusammenfassung der Änderungen", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Bereich", "Anzahl Änderungen", "Priorität"])

    summary = [
        ("Registrierungsprozess", "10 Änderungen", "HOCH"),
        ("Koordinator-Genehmigung", "5 neue Felder", "HOCH"),
        ("Datenmodell", "4 Modell-Änderungen, 2 neue Modelle", "HOCH"),
        ("Zeiterfassung", "3 Anpassungen", "MITTEL"),
        ("Admin-Funktionen", "3 neue Features", "MITTEL"),
        ("Dokumenten-Prozess", "2 Änderungen", "NIEDRIG"),
        ("Budget/Freibetrag", "2 Anpassungen", "MITTEL"),
    ]
    for bereich, anzahl, prio in summary:
        row = add_data_row(table, [bereich, anzahl, prio])
        prio_colors = {"HOCH": "F4CCCC", "MITTEL": "FFF2CC", "NIEDRIG": "D9EAD3"}
        set_cell_shading(row.cells[2], prio_colors.get(prio, "FFFFFF"))

    return doc


# ══════════════════════════════════════════════════════════════════════════
# DOCUMENT 2: IST_SOLL_VERGLEICH.docx — Compact side-by-side table
# ══════════════════════════════════════════════════════════════════════════

def create_ist_soll_vergleich():
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_heading("IST / SOLL Vergleich — Betreuer-App V2", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Stand: 03.03.2026 | CSFV Minden e.V.")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Legend
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Legende: ")
    run.bold = True
    p.add_run("GEÄNDERT = bestehender Prozess wird angepasst | NEU = komplett neuer Prozess | ENTFÄLLT = Prozess wird entfernt | UNVERÄNDERT = bleibt wie bisher")

    doc.add_paragraph()

    # ── MAIN TABLE: ALL CHANGES ───────────────────────────────────────

    all_rows = [
        # Section headers use None for IST/SOLL
        ("REGISTRIERUNG — Betreuer", None, None, None, None),
        ("1.1", "Registrierungslink", "Token-basiert, vom Koordinator erstellt, single/multi-use, mit Ablaufdatum", "Fester, permanenter Link /registrierung/ — immer erreichbar, kein Token nötig", "GEÄNDERT"),
        ("1.2", "Passwort", "System generiert Passwort", "Betreuer legt eigenes Passwort im Registrierungsformular fest", "GEÄNDERT"),
        ("1.3", "Registrierungsformular", "Persönliche Daten, Adresse, Bankdaten, Freibetrag-Erklärung", "Zusätzlich: Tätigkeitsart wählen + Schule wählen", "ERWEITERT"),
        ("1.4", "Duplikat-Prüfung", "Keine Prüfung auf Duplikate", "Hash-Prüfung (Vorname+Nachname+Geburtsdatum) — Duplikate werden erkannt und abgelehnt", "NEU"),
        ("1.5", "Email-Prüfung", "Kein Email-Abgleich", "Wenn Betreuer sich mit anderer Email registriert als bekannt → Hinweis an Betreuer", "NEU"),
        ("1.6", "Mehrfach-Anmeldung", "Ein Vertrag pro Betreuer pro Schuljahr", "Betreuer meldet sich für jedes Förderprogramm/Schule einzeln an — mehrere Verträge möglich", "GEÄNDERT"),
        ("1.7", "Führungszeugnis", "Erforderlich wenn is_external=True", "Erforderlich wenn Betreuer volljährig (>= 18 Jahre bei Anmeldung)", "GEÄNDERT"),
        ("1.8", "IBAN-Speicherung", "Fernet-verschlüsselt (EncryptedCharField)", "Unverschlüsselt (normales CharField)", "VEREINFACHT"),
        ("1.9", "Email-Pflicht", "Email optional", "Email ist Pflichtfeld für jeden Betreuer", "NEU"),

        ("GENEHMIGUNG — Koordinator", None, None, None, None),
        ("2.1", "Link-Verwaltung", "Koordinator erstellt & verwaltet Registrierungslinks", "Entfällt — Link ist permanent", "ENTFÄLLT"),
        ("2.2", "Anmeldungen prüfen", "Koordinator sieht Registrierungen seiner Schule(n)", "Unverändert", "UNVERÄNDERT"),
        ("2.3", "Genehmigung", "Koordinator genehmigt/lehnt ab", "Unverändert — Koordinator genehmigt jede Anmeldung", "UNVERÄNDERT"),
        ("2.4", "Förderprogramm zuordnen", "—", "NEU: Koordinator ordnet Förderprogramm zu (Standard wird vorgeschlagen)", "NEU"),
        ("2.5", "Vertragsbeginn", "Automatisch bei Vertragserstellung", "NEU: Koordinator legt Vertragsbeginn manuell fest", "NEU"),
        ("2.6", "Betreuer-Typ", "Wird bei Registrierung einmalig gesetzt", "NEU: Koordinator ergänzt/bestätigt Betreuer-Typ bei Genehmigung", "NEU"),
        ("2.7", "AG-Name", "ag_name-Feld im Vertrag (optional)", "NEU: Bei AG-Leitung muss Koordinator genauen AG-Namen eintragen", "NEU"),
        ("2.8", "Tätigkeit zuordnen", "Aus Vertrag übernommen", "NEU: Koordinator ordnet/bestätigt Tätigkeit bei Genehmigung zu", "NEU"),
        ("2.9", "Nach Genehmigung", "Dokumente werden generiert", "Email an Betreuer + Email an Personalabteilung (via N8N), dann Dokumente generiert", "UNVERÄNDERT"),

        ("ADMIN — Verwaltung", None, None, None, None),
        ("3.1", "Stammdaten", "Admin verwaltet Schulen, Schuljahre, FP, Kostenstellen, Stundensätze", "Unverändert", "UNVERÄNDERT"),
        ("3.2", "Übungsleiterpauschale", "Freibetrag-Limit als Feld in SchoolYear (3.300 EUR)", "§3 Nr. 26 EStG als eigenes Stammdatum — pro Kalenderjahr konfigurierbar", "GEÄNDERT"),
        ("3.3", "Manuelle Kosten", "—", "NEU: Admin kann manuell Kosten einem Förderprogramm zuordnen", "NEU"),
        ("3.4", "Auswertungen", "Keine Auswertung pro Förderprogramm", "NEU: Reports pro Förderprogramm — tatsächlich gezahlte Beträge", "NEU"),
        ("3.5", "Link-Verwaltung", "Admin verwaltet Registrierungslinks", "Entfällt — Link ist permanent", "ENTFÄLLT"),

        ("ZEITERFASSUNG — Betreuer", None, None, None, None),
        ("4.1", "Zeiterfassung", "Zeiten pro Vertrag erfassen", "Zeiten separat pro Schule erfassen", "GEÄNDERT"),
        ("4.2", "Genehmigung-Routing", "Alle Zeiten → ein Koordinator", "Je nach Schule → jeweiliger Koordinator", "GEÄNDERT"),
        ("4.3", "Förderprogramme", "Ein Vertrag pro FP", "Verschiedene Förderprogramme pro Schule möglich", "ERWEITERT"),
        ("4.4", "Zahlungsanweisung", "Eine Zahlungsanweisung nach Genehmigung", "Zahlungsanweisung pro Förderprogramm (KN, PN, Betrag)", "PRÄZISIERT"),

        ("DOKUMENTE", None, None, None, None),
        ("5.1", "Führungszeugnis-Regel", "Basiert auf is_external", "Basiert auf Alter >= 18 bei Anmeldung", "GEÄNDERT"),
        ("5.2", "Vertragsgenerierung", "Auto-generiert nach Registrierung", "Generiert nach Koordinator-Genehmigung (da KO Vertragsbeginn + FP ergänzt)", "GEÄNDERT"),
        ("5.3", "Sonstige Dokumente", "IfSB, Vertraulichkeit, Masernschutz", "Unverändert", "UNVERÄNDERT"),

        ("BUDGET & FREIBETRAG", None, None, None, None),
        ("6.1", "Freibetrag-Modell", "Feld in SchoolYear", "Eigenes Modell: Kalenderjahr + Betrag + §3 Nr. 26 EStG", "GEÄNDERT"),
        ("6.2", "Auszahlungen", "Nicht explizit KJ-bezogen", "Kalenderjahr-bezogen, an §3/26-Betrag geknüpft", "PRÄZISIERT"),
        ("6.3", "FP-Auswertung", "Kostenstelle hat FK zu FP", "Erweitert: tatsächlich gezahlte Beträge pro FP auswertbar", "ERWEITERT"),
    ]

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["#", "Bereich", "IST (Aktuell)", "SOLL (Neu)", "Status"])

    for row_data in all_rows:
        nr, bereich, ist, soll, status = row_data

        if ist is None:  # Section header
            row = table.add_row()
            # Merge all cells for section header
            row.cells[0].merge(row.cells[4])
            set_cell_shading(row.cells[0], "34495E")
            set_cell_text(row.cells[0], nr, bold=True, size=11, color=(255, 255, 255))
        else:
            add_data_row(table, [nr, bereich, ist, soll, status], status)

    # Set column widths
    for row in table.rows:
        if not row.cells[0]._element.getparent() is None:
            try:
                row.cells[0].width = Cm(1.0)
                row.cells[1].width = Cm(3.5)
                row.cells[2].width = Cm(9.0)
                row.cells[3].width = Cm(10.0)
                row.cells[4].width = Cm(3.0)
            except Exception:
                pass

    return doc


# ══════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    output_dir = os.path.dirname(os.path.abspath(__file__))

    # Document 1: Full process document
    doc1 = create_prozessaenderungen()
    path1 = os.path.join(output_dir, "PROZESSAENDERUNGEN_V2.docx")
    doc1.save(path1)
    print(f"[OK] {path1}")

    # Document 2: Compact IST/SOLL comparison
    doc2 = create_ist_soll_vergleich()
    path2 = os.path.join(output_dir, "IST_SOLL_VERGLEICH.docx")
    doc2.save(path2)
    print(f"[OK] {path2}")
