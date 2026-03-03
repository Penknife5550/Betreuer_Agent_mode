"""
Generate updated IST/SOLL Vergleich V2 Word document.
All open questions resolved. Includes N8N email notification matrix.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elem)


def set_cell_text(cell, text, bold=False, size=9, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_rich(cell, parts, size=9):
    """Set cell with mixed bold/normal text. parts = [(text, bold), ...]"""
    cell.text = ""
    p = cell.paragraphs[0]
    for text, bold in parts:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        if bold:
            run.bold = True


def add_header_row(table, texts, color_hex="2B579A"):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        set_cell_shading(cell, color_hex)
        set_cell_text(cell, text, bold=True, size=10, color=(255, 255, 255))


def add_data_row(table, texts, change_type=None):
    row = table.add_row()
    for i, text in enumerate(texts):
        set_cell_text(row.cells[i], text, size=9)
    if change_type:
        last_cell = row.cells[len(texts) - 1]
        colors = {
            "GEÄNDERT": "FFF2CC",
            "NEU": "D9EAD3",
            "ENTFÄLLT": "F4CCCC",
            "ERWEITERT": "D0E0F0",
            "VEREINFACHT": "E8D5F5",
            "PRÄZISIERT": "D0E0F0",
            "UNVERÄNDERT": "EEEEEE",
            "GEKLÄRT": "C8E6C9",
        }
        bg = colors.get(change_type, "FFFFFF")
        set_cell_shading(last_cell, bg)
        set_cell_text(last_cell, change_type, bold=True, size=9)
    return row


def add_section_header(table, text):
    row = table.add_row()
    row.cells[0].merge(row.cells[len(row.cells) - 1])
    set_cell_shading(row.cells[0], "34495E")
    set_cell_text(row.cells[0], text, bold=True, size=11, color=(255, 255, 255))
    return row


def set_col_widths(table, widths_cm):
    for row in table.rows:
        try:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)
        except Exception:
            pass


def create_document():
    doc = Document()

    # ── Page setup: A4 landscape ──
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading("IST / SOLL Vergleich", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Betreuer-App V2 — CSFV Minden e.V.", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Stand: 03.03.2026 | Alle offenen Fragen geklärt")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    # Decision summary box
    doc.add_heading("Getroffene Entscheidungen", level=2)
    decisions = [
        ("Hash-Bildung:", "Vorname + Nachname + Geburtsdatum (SHA256)"),
        ("Datenimport:", "Kein Import — kompletter Neustart"),
        ("Mehrfach-Anmeldung:", "Gespeicherte Daten werden wiederverwendet (vereinfachter Prozess)"),
        ("Gesamtschule/Gymnasium:", "Kombinierte Option in der Schulauswahl"),
        ("Vertragsende:", "Automatisch = Schuljahresende"),
        ("Onboarding-Status:", "Neuer Status 'pending_approval' wird eingeführt"),
        ("N8N-Emails:", "Alle Prozessschritte werden per N8N-Email benachrichtigt"),
        ("Formulare:", "Links zu MS Forms noch ausstehend (404) — Felder werden nachgetragen"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    add_header_row(table, ["Thema", "Entscheidung"], "1B5E20")
    for thema, entscheidung in decisions:
        row = table.add_row()
        set_cell_text(row.cells[0], thema, bold=True, size=10)
        set_cell_text(row.cells[1], entscheidung, size=10)
    set_col_widths(table, [6.0, 20.5])

    doc.add_paragraph()

    # Legend
    p = doc.add_paragraph()
    run = p.add_run("Legende Status-Spalte: ")
    run.bold = True
    run.font.size = Pt(9)
    p.add_run("GEÄNDERT = angepasst | NEU = komplett neu | ENTFÄLLT = wird entfernt | ERWEITERT = erweiterte Funktion | VEREINFACHT = reduziert | UNVERÄNDERT = bleibt").font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # HAUPTTABELLE: IST / SOLL
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("1. Prozessvergleich IST / SOLL", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["#", "Bereich", "IST (Aktuell)", "SOLL (Neu)", "Status"])

    # ── REGISTRIERUNG ─────────────────────────────────────────────
    add_section_header(table, "REGISTRIERUNG — Betreuer")

    rows = [
        ("1.1", "Registrierungslink",
         "Token-basiert, vom Koordinator erstellt, single/multi-use, mit Ablaufdatum",
         "Fester, permanenter Link /registrierung/ — immer erreichbar, kein Token nötig. Kein Koordinator-Eingriff mehr nötig.",
         "GEÄNDERT"),
        ("1.2", "Passwort",
         "System generiert Passwort",
         "Betreuer legt eigenes Passwort im Registrierungsformular fest",
         "GEÄNDERT"),
        ("1.3", "Formularfelder",
         "Persönliche Daten, Adresse, Bankdaten (IBAN verschlüsselt), Freibetrag-Erklärung",
         "Zusätzlich: Tätigkeitsart wählen + Schule wählen. Felder gem. FES-Personal-Betreuerfragebogen (MS Forms-Link ausstehend)",
         "ERWEITERT"),
        ("1.4", "Schulauswahl",
         "Schule wird über den Registrierungslink vorgegeben",
         "Betreuer wählt selbst: Gesamtschule/Gymnasium (kombiniert), GS Haddenhausen, GS Minderheide, GS Stemwede",
         "GEÄNDERT"),
        ("1.5", "Tätigkeitsart",
         "Wird im Vertrag festgelegt (nicht bei Registrierung)",
         "Betreuer wählt bei Registrierung: AG-Leitung, Hausaufgabenbetreuung, Hausaufgabenhilfe plus, Aufsicht, Päd. Assistenz, Schwimmbegleitung",
         "NEU"),
        ("1.6", "Duplikat-Prüfung",
         "Keine Prüfung auf Duplikate",
         "SHA256-Hash aus Vorname + Nachname + Geburtsdatum. Bei Match: Daten des bestehenden Profils werden wiederverwendet, nur neue Schule/Tätigkeit wird ergänzt.",
         "NEU"),
        ("1.7", "Email-Prüfung",
         "Kein Email-Abgleich (kein Import-Bestand)",
         "Bei Hash-Match mit bestehendem Profil: Wenn neue Email ≠ gespeicherte Email → Hinweis an Betreuer. Kein Import nötig (kompletter Neustart).",
         "NEU"),
        ("1.8", "Mehrfach-Anmeldung",
         "Ein Vertrag pro Betreuer pro Schuljahr",
         "Betreuer kann sich mehrfach anmelden (verschiedene Schulen/FP). Bei Wiedererkennung (Hash-Match) werden persönliche Daten wiederverwendet — nur Schule + Tätigkeit werden neu abgefragt.",
         "GEÄNDERT"),
        ("1.9", "Führungszeugnis",
         "Erforderlich wenn is_external=True",
         "Erforderlich wenn Betreuer volljährig (>= 18 Jahre bei Anmeldung). Minderjährige sind befreit.",
         "GEÄNDERT"),
        ("1.10", "IBAN-Speicherung",
         "Fernet-verschlüsselt (EncryptedCharField)",
         "Unverschlüsselt (normales CharField) — kein Fernet mehr nötig",
         "VEREINFACHT"),
        ("1.11", "Email-Pflicht",
         "Email optional",
         "Email ist Pflichtfeld. Wird für alle N8N-Benachrichtigungen benötigt.",
         "NEU"),
        ("1.12", "N8N-Email",
         "Event 'betreuer_registered' an N8N",
         "Nach Registrierung: Email an Koordinator der gewählten Schule ('neue Anmeldung prüfen') via N8N",
         "ERWEITERT"),
    ]
    for r in rows:
        add_data_row(table, list(r), r[4])

    # ── GENEHMIGUNG ───────────────────────────────────────────────
    add_section_header(table, "GENEHMIGUNG — Koordinator")

    rows = [
        ("2.1", "Link-Verwaltung",
         "Koordinator erstellt & verwaltet Registrierungslinks",
         "Entfällt komplett — Link ist permanent, kein Koordinator-Eingriff",
         "ENTFÄLLT"),
        ("2.2", "Anmeldungen sehen",
         "Koordinator sieht Registrierungen seiner Schule(n)",
         "Unverändert — Anmeldungen werden nach Schule geroutet",
         "UNVERÄNDERT"),
        ("2.3", "Genehmigung",
         "Koordinator genehmigt/lehnt ab",
         "Unverändert — Koordinator genehmigt jede Anmeldung. Neuer Onboarding-Status: registered → pending_approval → approved",
         "ERWEITERT"),
        ("2.4", "Förderprogramm",
         "—",
         "NEU: Koordinator ordnet Förderprogramm zu. Standard wird automatisch vorgeschlagen: Weiterführend → 'Geld oder Stelle', GS → 'Schule von 8 bis 1', AG an GS → '13 plus'",
         "NEU"),
        ("2.5", "Vertragsbeginn",
         "Automatisch bei Vertragserstellung",
         "NEU: Koordinator legt Vertragsbeginn manuell fest. Vertragsende = automatisch Schuljahresende.",
         "NEU"),
        ("2.6", "Betreuer-Typ",
         "Wird bei Registrierung einmalig gesetzt",
         "NEU: Koordinator ergänzt Betreuer-Typ bei Genehmigung (Schüler, sonst. MA, langjährig, Lehrer, LA-Student, extern). Felder gem. Koordinator-Fragebogen (MS Forms-Link ausstehend).",
         "NEU"),
        ("2.7", "AG-Name",
         "ag_name-Feld im Vertrag (optional)",
         "NEU: Bei Tätigkeit 'AG-Leitung' muss Koordinator genauen AG-Namen eintragen (Pflichtfeld)",
         "NEU"),
        ("2.8", "Tätigkeit zuordnen",
         "Tätigkeit aus Vertrag übernommen",
         "NEU: Koordinator bestätigt/ändert die vom Betreuer gewählte Tätigkeit",
         "NEU"),
        ("2.9", "N8N nach Genehmigung",
         "Event 'betreuer_registered' → Dokumente generiert",
         "Nach Genehmigung via N8N: (1) Email an Betreuer ('Anmeldung genehmigt'), (2) Email an Personalabteilung ('neuer Betreuer'), (3) System generiert Dokumente",
         "ERWEITERT"),
        ("2.10", "N8N bei Ablehnung",
         "—",
         "NEU: Bei Ablehnung durch Koordinator → Email an Betreuer via N8N mit Ablehnungsgrund",
         "NEU"),
    ]
    for r in rows:
        add_data_row(table, list(r), r[4])

    # ── ADMIN ─────────────────────────────────────────────────────
    add_section_header(table, "ADMIN — Verwaltung & Stammdaten")

    rows = [
        ("3.1", "Stammdaten",
         "Admin verwaltet Schulen, Schuljahre, FP, Kostenstellen, Stundensätze",
         "Unverändert",
         "UNVERÄNDERT"),
        ("3.2", "Übungsleiterpauschale",
         "Freibetrag-Limit als Feld in SchoolYear (default 3.300 EUR, Schuljahr-bezogen)",
         "Neues Modell: §3 Nr. 26 EStG Übungsleiterpauschale — pro Kalenderjahr konfigurierbar (Betrag ändert sich gesetzlich). Auszahlungsbudget an diesen Betrag geknüpft.",
         "GEÄNDERT"),
        ("3.3", "Manuelle Kosten",
         "—",
         "NEU: Admin kann manuell Kosten einem Förderprogramm zuordnen (z.B. Material, Fortbildung). Getrennt von Stundennachweis-Abrechnungen.",
         "NEU"),
        ("3.4", "Auswertungen/Reports",
         "Monatliche Übersicht + Freibetrag-Übersicht",
         "NEU: Zusätzlich Reports pro Förderprogramm — tatsächlich gezahlte Beträge (über Kostenstelle → FP). CSV-Export.",
         "NEU"),
        ("3.5", "Link-Verwaltung",
         "Admin verwaltet Registrierungslinks",
         "Entfällt — Link ist permanent",
         "ENTFÄLLT"),
        ("3.6", "Betreuer aktivieren",
         "Admin aktiviert Betreuer (Status → aktiv)",
         "Unverändert — nach Dokumenten-Prüfung",
         "UNVERÄNDERT"),
    ]
    for r in rows:
        add_data_row(table, list(r), r[4])

    # ── ZEITERFASSUNG ─────────────────────────────────────────────
    add_section_header(table, "ZEITERFASSUNG & ABRECHNUNG")

    rows = [
        ("4.1", "Zeiterfassung",
         "Betreuer erfasst Zeiten pro Vertrag (eine Schule)",
         "Betreuer erfasst Zeiten separat pro Schule. Dashboard zeigt alle Schulen, Betreuer wählt Kontext.",
         "GEÄNDERT"),
        ("4.2", "Genehmigung-Routing",
         "Alle Zeiten gehen an einen Koordinator",
         "Je nach Schule → Genehmigung geht an den jeweiligen Schulkoordinator",
         "GEÄNDERT"),
        ("4.3", "Förderprogramme",
         "Ein Vertrag, ein Förderprogramm",
         "Betreuer kann an verschiedenen Förderprogrammen pro Schule angemeldet sein. Zeiterfassung wird dem richtigen FP zugeordnet.",
         "ERWEITERT"),
        ("4.4", "Zahlungsanweisung",
         "Eine Zahlungsanweisung nach Genehmigung",
         "Zahlungsanweisung pro Förderprogramm nach Genehmigung (Kreditorennr., Projektnr., Betrag als QR-Code)",
         "PRÄZISIERT"),
        ("4.5", "Stundennachweise",
         "Koordinator sieht Stundennachweise seiner Schule(n)",
         "Unverändert",
         "UNVERÄNDERT"),
        ("4.6", "N8N nach Einreichung",
         "Event 'timesheet_submitted'",
         "Email an Koordinator via N8N: 'Neuer Stundennachweis zur Prüfung' mit Details (Betreuer, Monat, Stunden)",
         "ERWEITERT"),
        ("4.7", "N8N nach Genehmigung",
         "Event 'timesheet_approved' → N8N an Buchhaltung",
         "Erweitert: (1) Email an Betreuer ('Stunden genehmigt'), (2) Email an Buchhaltung (Zahlungsanweisung PDF, KN, PN), (3) Freibetrag-Warnung wenn Schwelle erreicht",
         "ERWEITERT"),
        ("4.8", "N8N nach Ablehnung",
         "Betreuer wird benachrichtigt",
         "Email an Betreuer via N8N: 'Stunden abgelehnt' mit Ablehnungsgrund. Betreuer kann korrigieren & erneut einreichen.",
         "ERWEITERT"),
    ]
    for r in rows:
        add_data_row(table, list(r), r[4])

    # ── DOKUMENTE ─────────────────────────────────────────────────
    add_section_header(table, "DOKUMENTE")

    rows = [
        ("5.1", "Führungszeugnis-Regel",
         "Erforderlich wenn is_external=True",
         "Erforderlich wenn Betreuer volljährig (>= 18 Jahre bei Anmeldung). Minderjährige befreit.",
         "GEÄNDERT"),
        ("5.2", "Vertragsgenerierung",
         "Auto-generiert direkt nach Registrierung",
         "Generiert erst nach Koordinator-Genehmigung (da Koordinator Vertragsbeginn, FP, Betreuer-Typ ergänzt). Vertragsende = Schuljahresende (automatisch).",
         "GEÄNDERT"),
        ("5.3", "Vertraulichkeitserklärung",
         "Auto-generiert",
         "Unverändert — nach Koordinator-Genehmigung generiert",
         "UNVERÄNDERT"),
        ("5.4", "IfSB",
         "Auto-generiert, 24 Monate Renewal",
         "Unverändert",
         "UNVERÄNDERT"),
        ("5.5", "Masernschutz",
         "Manueller Upload durch Betreuer",
         "Unverändert",
         "UNVERÄNDERT"),
        ("5.6", "N8N Dokumente",
         "Events: document_generated, document_expiring, document_expired",
         "Erweitert: (1) Email an Betreuer wenn Dokumente zum Upload bereit, (2) Erinnerung bei ablaufenden Dokumenten, (3) Email an Koordinator wenn alle Dokumente verifiziert",
         "ERWEITERT"),
    ]
    for r in rows:
        add_data_row(table, list(r), r[4])

    # ── BUDGET & FREIBETRAG ───────────────────────────────────────
    add_section_header(table, "BUDGET & FREIBETRAG")

    rows = [
        ("6.1", "Freibetrag-Modell",
         "Feld 'freibetrag_limit' in SchoolYear (Schuljahr-bezogen)",
         "Eigenes Modell 'Uebungsleiterpauschale': Kalenderjahr + Betrag + §3 Nr. 26 EStG. Saubere Trennung Schuljahr vs. Kalenderjahr.",
         "GEÄNDERT"),
        ("6.2", "FP-Budget",
         "Budget-Feld im Förderprogramm (Schuljahr-bezogen)",
         "Unverändert — Budget bleibt Schuljahr-bezogen",
         "UNVERÄNDERT"),
        ("6.3", "Auszahlungen",
         "Nicht explizit Kalenderjahr-bezogen",
         "Auszahlungen sind Kalenderjahr-bezogen. Budget für Betreuer-Auszahlungen ist an den §3/26-Betrag geknüpft.",
         "PRÄZISIERT"),
        ("6.4", "FP-Auswertung",
         "Kostenstelle hat FK zu Förderprogramm (vorhanden)",
         "Erweitert: Auswertung tatsächlich gezahlter Beträge pro FP (genehmigte Stundennachweise + manuelle Kosten)",
         "ERWEITERT"),
        ("6.5", "N8N Freibetrag",
         "Event 'freibetrag_warning' bei 80/90/100%",
         "Erweitert: (1) Email an Betreuer bei Schwelle, (2) Email an Admin/Koordinator bei 90%+, (3) Sperre bei 100% mit Email an alle Beteiligten",
         "ERWEITERT"),
    ]
    for r in rows:
        add_data_row(table, list(r), r[4])

    set_col_widths(table, [1.0, 3.5, 9.0, 10.0, 2.8])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # ONBOARDING-STATUS NEU
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("2. Onboarding-Status (NEU vs. ALT)", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["Status ALT", "Status NEU", "Beschreibung", "N8N-Email"])

    status_rows = [
        ("registered", "registered", "Betreuer hat Registrierung abgeschlossen", "→ Email an Koordinator: 'Neue Anmeldung prüfen'"),
        ("—", "pending_approval (NEU)", "Wartet auf Genehmigung durch Koordinator. Koordinator ergänzt: FP, Vertragsbeginn, Betreuer-Typ, ggf. AG-Name", "—"),
        ("—", "approved (NEU)", "Koordinator hat genehmigt. System generiert Vertrag + Dokumente", "→ Email an Betreuer: 'Genehmigt'\n→ Email an Personalabt.: 'Neuer Betreuer'"),
        ("documents_pending", "documents_pending", "Dokumente generiert, Betreuer muss signiert hochladen", "→ Email an Betreuer: 'Dokumente bereit zum Upload'"),
        ("documents_complete", "documents_complete", "Alle Dokumente verifiziert durch Koordinator", "→ Email an Admin: 'Betreuer bereit zur Aktivierung'\n→ Email an Koordinator: 'Dokumente vollständig'"),
        ("active", "active", "Admin hat Betreuer aktiviert — kann Zeiten erfassen", "→ Email an Betreuer: 'Konto aktiviert'"),
        ("inactive", "inactive", "Betreuer deaktiviert (z.B. Schuljahresende)", "→ Email an Betreuer: 'Konto deaktiviert'"),
        ("archived", "archived", "Betreuer archiviert", "—"),
    ]
    for alt, neu, desc, email in status_rows:
        row = add_data_row(table, [alt, neu, desc, email])
        if "NEU" in neu:
            set_cell_shading(row.cells[1], "D9EAD3")
            set_cell_text(row.cells[1], neu, bold=True, size=9)
    set_col_widths(table, [3.5, 4.0, 9.0, 10.0])

    doc.add_paragraph()

    # Status flow diagram as text
    p = doc.add_paragraph()
    run = p.add_run("Status-Flow NEU: ")
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(
        "registered → pending_approval → approved → documents_pending → documents_complete → active → inactive → archived"
    )
    p = doc.add_paragraph()
    run = p.add_run("Ablehnung: ")
    run.bold = True
    p.add_run("pending_approval → rejected → (Betreuer meldet sich erneut an) → registered")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # N8N EMAIL-MATRIX
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("3. N8N Email-Benachrichtigungen (Komplett)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Alle Prozessschritte werden über N8N per Email kommuniziert.")
    run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["#", "Event / Auslöser", "Empfänger", "Email-Inhalt (Kurzfassung)", "Prozess"])

    n8n_rows = [
        ("E1", "Neue Registrierung", "Koordinator der gewählten Schule", "Neue Betreuer-Anmeldung eingegangen. Bitte prüfen und genehmigen.", "Registrierung"),
        ("E2", "Anmeldung genehmigt", "Betreuer", "Ihre Anmeldung als [Tätigkeit] an [Schule] wurde genehmigt. Bitte laden Sie die Dokumente hoch.", "Genehmigung"),
        ("E3", "Anmeldung genehmigt", "Personalabteilung", "Neuer Betreuer [Name] für [Schule], FP: [Förderprogramm], Beginn: [Datum].", "Genehmigung"),
        ("E4", "Anmeldung abgelehnt", "Betreuer", "Ihre Anmeldung wurde abgelehnt. Grund: [Ablehnungsgrund].", "Genehmigung"),
        ("E5", "Dokumente generiert", "Betreuer", "Ihre Dokumente (Vertrag, Vertraulichkeit, IfSB) stehen zum Download/Upload bereit.", "Dokumente"),
        ("E6", "Dokument hochgeladen", "Koordinator", "Betreuer [Name] hat [Dokumenttyp] hochgeladen. Bitte prüfen.", "Dokumente"),
        ("E7", "Dokument verifiziert", "Betreuer", "Ihr Dokument [Typ] wurde geprüft und akzeptiert.", "Dokumente"),
        ("E8", "Dokument abgelehnt", "Betreuer", "Ihr Dokument [Typ] wurde abgelehnt. Grund: [Grund]. Bitte erneut hochladen.", "Dokumente"),
        ("E9", "Alle Dokumente komplett", "Admin + Koordinator", "Betreuer [Name] — alle Dokumente vollständig. Bereit zur Aktivierung.", "Dokumente"),
        ("E10", "Betreuer aktiviert", "Betreuer", "Ihr Konto wurde aktiviert. Sie können ab sofort Ihre Zeiten erfassen.", "Aktivierung"),
        ("E11", "Stundennachweis eingereicht", "Koordinator der Schule", "Neuer Stundennachweis von [Name] für [Monat/Jahr]. Bitte prüfen und genehmigen.", "Zeiterfassung"),
        ("E12", "Stundennachweis genehmigt", "Betreuer", "Ihr Stundennachweis für [Monat/Jahr] wurde genehmigt. Betrag: [EUR].", "Zeiterfassung"),
        ("E13", "Stundennachweis genehmigt", "Buchhaltung", "Zahlungsanweisung für [Name], FP: [FP], KN: [KN], PN: [PN], Betrag: [EUR]. PDF im Anhang.", "Zeiterfassung"),
        ("E14", "Stundennachweis abgelehnt", "Betreuer", "Ihr Stundennachweis für [Monat/Jahr] wurde abgelehnt. Grund: [Grund]. Bitte korrigieren.", "Zeiterfassung"),
        ("E15", "Freibetrag 80% erreicht", "Betreuer", "Hinweis: Sie haben 80% Ihres Freibetrags (§3 Nr. 26 EStG) aufgebraucht.", "Freibetrag"),
        ("E16", "Freibetrag 90% erreicht", "Betreuer + Koordinator + Admin", "Warnung: Betreuer [Name] hat 90% des Freibetrags erreicht.", "Freibetrag"),
        ("E17", "Freibetrag 100% erreicht", "Betreuer + Koordinator + Admin", "ACHTUNG: Freibetrag [Name] ausgeschöpft. Weitere Auszahlungen wären steuerpflichtig.", "Freibetrag"),
        ("E18", "Dokument läuft ab (30 Tage)", "Betreuer + Koordinator", "Dokument [Typ] von [Name] läuft in 30 Tagen ab. Bitte erneuern.", "Dokumente"),
        ("E19", "Dokument abgelaufen", "Betreuer + Koordinator + Admin", "Dokument [Typ] von [Name] ist abgelaufen.", "Dokumente"),
    ]
    for nr, event, empf, inhalt, prozess in n8n_rows:
        row = add_data_row(table, [nr, event, empf, inhalt, prozess])
        # Color by process
        proc_colors = {
            "Registrierung": "D0E0F0",
            "Genehmigung": "D9EAD3",
            "Dokumente": "FFF2CC",
            "Aktivierung": "C8E6C9",
            "Zeiterfassung": "E8D5F5",
            "Freibetrag": "F4CCCC",
        }
        set_cell_shading(row.cells[4], proc_colors.get(prozess, "FFFFFF"))

    set_col_widths(table, [1.0, 4.0, 5.0, 12.5, 3.0])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # DATENMODELL-ÄNDERUNGEN
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("4. Datenmodell-Änderungen", level=1)

    # 4.1 BetreuerProfile
    doc.add_heading("4.1 BetreuerProfile", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    add_header_row(table, ["Feld", "IST", "SOLL", "Status"])

    bp_rows = [
        ("iban", "EncryptedCharField (Fernet)", "CharField (unverschlüsselt)", "GEÄNDERT"),
        ("unique_hash", "— (nicht vorhanden)", "CharField, SHA256 aus Vorname+Nachname+Geburtsdatum, unique=True", "NEU"),
        ("email (User.email)", "Optional", "Pflichtfeld (required=True)", "GEÄNDERT"),
        ("requires_fuehrungszeugnis", "Property: self.is_external", "Property: Alter >= 18 (basierend auf geburtsdatum bei Anmeldung)", "GEÄNDERT"),
        ("onboarding_status", "Choices: registered, documents_pending, documents_complete, active, inactive, archived", "NEU: + pending_approval, approved. Flow: registered → pending_approval → approved → documents_pending → documents_complete → active", "ERWEITERT"),
    ]
    for feld, ist, soll, status in bp_rows:
        add_data_row(table, [feld, ist, soll, status], status)
    set_col_widths(table, [4.0, 8.0, 11.0, 3.0])

    doc.add_paragraph()

    # 4.2 Contract
    doc.add_heading("4.2 Contract", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    add_header_row(table, ["Feld", "IST", "SOLL", "Status"])

    ct_rows = [
        ("start_date", "Automatisch bei Erstellung", "Vom Koordinator bei Genehmigung gesetzt", "GEÄNDERT"),
        ("end_date", "Manuell oder Schuljahresende", "Automatisch = Schuljahresende (berechnet aus aktuellem SchoolYear.end_date)", "GEÄNDERT"),
        ("foerderprogramme", "M2M, optional, manuell zugeordnet", "Vom Koordinator bei Genehmigung zugeordnet. Standard automatisch vorgeschlagen.", "GEÄNDERT"),
    ]
    for feld, ist, soll, status in ct_rows:
        add_data_row(table, [feld, ist, soll, status], status)
    set_col_widths(table, [4.0, 8.0, 11.0, 3.0])

    doc.add_paragraph()

    # 4.3 RegistrationLink
    doc.add_heading("4.3 RegistrationLink → ENTFÄLLT", level=2)
    doc.add_paragraph("Token-basierte Links werden komplett entfernt. Ersetzt durch festen Endpunkt /registrierung/.")

    doc.add_paragraph()

    # 4.4 School
    doc.add_heading("4.4 School — Anpassung Schulauswahl", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Anzeigename (Registrierung)", "Schultyp(en)", "Koordinator"])

    school_rows = [
        ("Gesamtschule/Gymnasium", "gesamtschule + gymnasium (kombiniert)", "Ein gemeinsamer Koordinator"),
        ("Grundschule Haddenhausen", "grundschule", "Eigener Koordinator"),
        ("Grundschule Minderheide", "grundschule", "Eigener Koordinator"),
        ("Grundschule Stemwede", "grundschule", "Eigener Koordinator"),
    ]
    for s, t, k in school_rows:
        add_data_row(table, [s, t, k])
    set_col_widths(table, [7.0, 8.0, 7.0])

    doc.add_paragraph()

    # 4.5 ActivityType
    doc.add_heading("4.5 ActivityType — Finale Liste", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    add_header_row(table, ["Code", "Tätigkeitsart", "Weiterführend (Geld/Stelle)", "Grundschule (8-1 / 13+)"])

    at_rows = [
        ("ag_leitung", "AG-Leitung", "Ja (Geld oder Stelle)", "Ja (13 plus)"),
        ("hausaufgabenbetreuung", "Hausaufgabenbetreuung", "Ja (Geld oder Stelle)", "Ja (Schule von 8 bis 1)"),
        ("hausaufgabenhilfe_plus", "Hausaufgabenhilfe plus", "Ja (Geld oder Stelle)", "Nein"),
        ("aufsicht", "Aufsicht", "Nein", "Ja (Schule von 8 bis 1)"),
        ("paed_assistenz", "Pädagogische Assistenz", "Nein", "Ja (Schule von 8 bis 1)"),
        ("schwimmbegleitung", "Schwimmbegleitung", "Ja (Geld oder Stelle)", "Nein"),
    ]
    for code, name, wf, gs in at_rows:
        row = add_data_row(table, [code, name, wf, gs])
        if wf.startswith("Nein"):
            set_cell_shading(row.cells[2], "F4CCCC")
        if gs.startswith("Nein"):
            set_cell_shading(row.cells[3], "F4CCCC")
    set_col_widths(table, [4.5, 5.0, 6.5, 6.5])

    doc.add_paragraph()

    # 4.6 Neue Modelle
    doc.add_heading("4.6 Uebungsleiterpauschale — NEUES MODELL", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Feld", "Typ", "Beschreibung"])
    for feld, typ, desc in [
        ("kalenderjahr", "IntegerField, unique", "z.B. 2026"),
        ("betrag", "DecimalField(10,2)", "z.B. 3.300,00 EUR"),
        ("gesetzliche_grundlage", "CharField", "§3 Nr. 26 EStG"),
        ("gueltig_ab", "DateField", "Ab wann dieser Betrag gilt"),
    ]:
        add_data_row(table, [feld, typ, desc])

    doc.add_paragraph()

    doc.add_heading("4.7 ManuelleKostenbuchung — NEUES MODELL", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Feld", "Typ", "Beschreibung"])
    for feld, typ, desc in [
        ("foerderprogramm", "ForeignKey(Foerderprogramm)", "Zuordnung zum Förderprogramm"),
        ("betrag", "DecimalField(10,2)", "Manuell eingetragener Betrag"),
        ("beschreibung", "TextField", "Freitext-Beschreibung"),
        ("kategorie", "CharField (choices)", "Material, Fortbildung, Versicherung, Sonstiges"),
        ("beleg_nr", "CharField (optional)", "Belegnummer für Buchhaltung"),
        ("datum", "DateField", "Buchungsdatum"),
        ("erstellt_von", "ForeignKey(User)", "Admin der die Buchung erstellt hat"),
    ]:
        add_data_row(table, [feld, typ, desc])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # FÖRDERPROGRAMM-ZUORDNUNGSLOGIK
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("5. Automatische Förderprogramm-Zuordnung", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Regelwerk für die automatische Vorauswahl des Förderprogramms bei der Koordinator-Genehmigung:")
    run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    add_header_row(table, ["Schule", "Tätigkeit", "→ Förderprogramm (Vorschlag)", "Koordinator"])

    logic_rows = [
        ("Gesamtschule/Gymnasium", "AG-Leitung", "Geld oder Stelle", "Kann ändern"),
        ("Gesamtschule/Gymnasium", "Hausaufgabenbetreuung", "Geld oder Stelle", "Kann ändern"),
        ("Gesamtschule/Gymnasium", "Hausaufgabenhilfe plus", "Geld oder Stelle", "Kann ändern"),
        ("Gesamtschule/Gymnasium", "Schwimmbegleitung", "Geld oder Stelle", "Kann ändern"),
        ("Grundschule (alle)", "AG-Leitung", "13 plus", "Kann ändern"),
        ("Grundschule (alle)", "Hausaufgabenbetreuung", "Schule von 8 bis 1", "Kann ändern"),
        ("Grundschule (alle)", "Pädagogische Assistenz", "Schule von 8 bis 1", "Kann ändern"),
        ("Grundschule (alle)", "Aufsicht", "Schule von 8 bis 1", "Kann ändern"),
    ]
    for s, t, fp, ko in logic_rows:
        row = add_data_row(table, [s, t, fp, ko])
        fp_colors = {"Geld oder Stelle": "D0E0F0", "13 plus": "FFF2CC", "Schule von 8 bis 1": "D9EAD3"}
        set_cell_shading(row.cells[2], fp_colors.get(fp, "FFFFFF"))
    set_col_widths(table, [5.5, 5.5, 6.0, 3.0])

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # ZUSAMMENFASSUNG
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("6. Zusammenfassung aller Änderungen", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    add_header_row(table, ["Bereich", "Änderungen", "Neue N8N-Events", "Priorität"])

    summary = [
        ("Registrierung", "12 (5 geändert, 4 neu, 1 erweitert, 1 vereinfacht, 1 entfällt)", "E1", "HOCH"),
        ("Koordinator-Genehmigung", "10 (5 neu, 2 erweitert, 1 entfällt, 2 unverändert)", "E2, E3, E4", "HOCH"),
        ("Admin-Verwaltung", "6 (1 geändert, 2 neu, 1 entfällt, 2 unverändert)", "—", "MITTEL"),
        ("Zeiterfassung", "8 (2 geändert, 1 erweitert, 3 N8N erweitert, 1 präzisiert, 1 unverändert)", "E11-E14", "HOCH"),
        ("Dokumente", "6 (2 geändert, 1 erweitert, 3 unverändert)", "E5-E9, E18, E19", "MITTEL"),
        ("Budget/Freibetrag", "5 (1 geändert, 1 präzisiert, 2 erweitert, 1 unverändert)", "E15-E17", "MITTEL"),
        ("Onboarding-Status", "2 neue Status (pending_approval, approved)", "Alle Status-Emails", "HOCH"),
        ("Datenmodell", "3 geänderte Modelle, 2 neue Modelle, 1 entfernt", "—", "HOCH"),
    ]
    for bereich, aend, events, prio in summary:
        row = add_data_row(table, [bereich, aend, events, prio])
        prio_colors = {"HOCH": "F4CCCC", "MITTEL": "FFF2CC", "NIEDRIG": "D9EAD3"}
        set_cell_shading(row.cells[3], prio_colors.get(prio, "FFFFFF"))
    set_col_widths(table, [4.0, 12.0, 4.0, 2.5])

    doc.add_paragraph()

    # Total count
    p = doc.add_paragraph()
    run = p.add_run("Gesamt: 47 Änderungspunkte | 19 N8N-Email-Events | 2 neue Modelle | 2 neue Onboarding-Status")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Offener Punkt: ")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run("Links zu den MS Forms-Formularen (Betreuer-Fragebogen + Koordinator-Fragebogen) konnten nicht abgerufen werden (404). Bitte korrekte Links nachliefern, damit die Formularfelder ins System übernommen werden können.")

    return doc


if __name__ == "__main__":
    output_dir = os.path.dirname(os.path.abspath(__file__))
    doc = create_document()
    path = os.path.join(output_dir, "IST_SOLL_VERGLEICH_V2.docx")
    doc.save(path)
    print(f"[OK] {path}")
